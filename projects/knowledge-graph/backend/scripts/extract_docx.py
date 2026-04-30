"""提取 AIXC Yellowpaper v6.2 中文结构 + 全文到 temp 文件。

打印 2 项：
1. 目录结构（所有 Heading 级段落）
2. 正文前 N 段的第一行，了解风格
然后把全文写到 workspace 的 tmp 位置供 Read 工具读取。
"""
import os, sys, io, json, re
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding="utf-8")
from docx import Document

src = os.path.join(os.environ["TEMP"], "aixc_yp.docx")
d = Document(src)

print(f"=== AIXC Yellowpaper v6.2 概况 ===")
print(f"段落数: {len(d.paragraphs)}")
print(f"表格数: {len(d.tables)}")
print()

print("=== 目录结构（标题级段落）===")
toc = []
for p in d.paragraphs:
    if not p.text.strip():
        continue
    style = p.style.name
    if style.startswith("Heading") or style in ("Title", "Subtitle"):
        toc.append((style, p.text))
        print(f"  [{style}] {p.text[:150]}")

print()
print(f"=== 目录共 {len(toc)} 条 ===")
print()

# 把全文导出成 md 便于 Read 工具读
out_md = r"C:\Users\ken\OneDrive\Document\Github\ai-social-research\.tmp_aixc_yp_v6.2.md"
os.makedirs(os.path.dirname(out_md), exist_ok=True)
with open(out_md, "w", encoding="utf-8") as f:
    f.write("# AIXC Yellowpaper v6.2 (提取自 docx)\n\n")
    for p in d.paragraphs:
        style = p.style.name
        text = p.text
        if not text.strip():
            f.write("\n")
            continue
        # 标题级别转 markdown
        if style.startswith("Heading"):
            try:
                lvl = int(style.replace("Heading ", "")) + 1
            except ValueError:
                lvl = 2
            f.write(f"{'#' * lvl} {text}\n\n")
        elif style in ("Title",):
            f.write(f"# {text}\n\n")
        elif style in ("Subtitle",):
            f.write(f"## {text}\n\n")
        else:
            f.write(f"{text}\n\n")
    # 表格
    for i, tbl in enumerate(d.tables):
        f.write(f"\n---\n\n## 表格 {i+1}\n\n")
        for row in tbl.rows:
            f.write("| " + " | ".join(c.text.replace("\n", " ").strip() for c in row.cells) + " |\n")
        f.write("\n")

print(f"全文导出到: {out_md}")
print(f"文件大小: {os.path.getsize(out_md)} bytes")
